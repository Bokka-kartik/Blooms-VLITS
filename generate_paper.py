
# import pandas as pd
# import random
# from docx import Document
# import os
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# DATASET_FOLDER = "dataset/"

# def get_dataset_for_subject(subject):
#     dataset_path = os.path.join(DATASET_FOLDER, f"{subject}.csv")
#     if os.path.exists(dataset_path):
#         return pd.read_csv(dataset_path)
#     return None

# def calculate_questions_per_bloom(num_questions, bloom_distribution):
#     questions_per_bloom = {}
#     remaining_questions = num_questions
    
#     # First pass - calculate whole numbers
#     for bloom, percentage in bloom_distribution.items():
#         if percentage == 0:
#             questions_per_bloom[bloom] = 0
#             continue
            
#         num = int((percentage / 100) * num_questions)
#         questions_per_bloom[bloom] = num
#         remaining_questions -= num
    
#     # Distribute remaining questions based on highest decimal parts
#     if remaining_questions > 0:
#         decimals = {
#             bloom: (percentage / 100 * num_questions) - questions_per_bloom[bloom]
#             for bloom, percentage in bloom_distribution.items()
#             if percentage > 0
#         }
        
#         sorted_blooms = sorted(decimals.items(), key=lambda x: x[1], reverse=True)
#         for i in range(remaining_questions):
#             if i < len(sorted_blooms):
#                 questions_per_bloom[sorted_blooms[i][0]] += 1
    
#     return questions_per_bloom

# def generate_question_paper(topic, num_questions, total_marks, bloom_distribution, difficulty_distribution, marks_per_difficulty):
#     # Initial validation
#     df = get_dataset_for_subject(topic)
#     if df is None or df.empty:
#         raise ValueError(f"No dataset found for topic: {topic}")
    
#     # Validate total marks calculation
#     expected_marks = sum(difficulty_distribution[level] * marks_per_difficulty[level] 
#                         for level in difficulty_distribution.keys())
#     if expected_marks != total_marks:
#         raise ValueError(f"Total marks calculation mismatch. Expected {total_marks}, got {expected_marks}")
    
#     # Validate difficulty distribution
#     if sum(difficulty_distribution.values()) != num_questions:
#         raise ValueError("Sum of difficulty distribution must equal total number of questions")
    
#     # Calculate number of questions needed for each Bloom's level
#     questions_per_bloom = calculate_questions_per_bloom(num_questions, bloom_distribution)
    
#     # Select questions based on Bloom's taxonomy and difficulty
#     selected_questions = []
#     questions_with_marks = []
    
#     grouped_questions = {
#         bloom: {
#            diff: df[(df['bloom_taxonomy_level'] == bloom) & 
#                     (df['difficulty'] == diff)]['question'].tolist()
#             for diff in difficulty_distribution.keys()
#         }
#         for bloom in questions_per_bloom.keys()
#     }


    
#     # Select questions ensuring both Bloom's and difficulty distributions are met
#     difficulty_counts = {diff: 0 for diff in difficulty_distribution.keys()}
    
#     for bloom, count in questions_per_bloom.items():
#         if count == 0:
#             continue
            
#         bloom_questions = []
#         for diff, target in difficulty_distribution.items():
#             available = grouped_questions[bloom][diff]
#             needed = min(
#                 target - difficulty_counts[diff],
#                 count - len(bloom_questions)
#             )
            
#             if needed <= 0:
#                 continue
                
#             if available:
#                 selected = random.sample(available, min(needed, len(available)))
#                 for q in selected:
#                     bloom_questions.append((q, diff))
#                     difficulty_counts[diff] += 1
                    
#             if len(bloom_questions) >= count:
#                 break
                
#         selected_questions.extend(bloom_questions)
    
#     # Shuffle questions to mix different types
#     random.shuffle(selected_questions)
    
#     # Create document with proper formatting
#     doc = Document()
    
#     # Add header
#     header = doc.add_heading(f"Question Paper: {topic}", level=1)
#     header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
#     # Add total marks
#     marks_para = doc.add_paragraph(f"Total Marks: {total_marks}")
#     marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
#     doc.add_paragraph("Time: 3 Hours") # You can make this configurable
#     doc.add_paragraph("")  # Spacing
    
#     # Add instructions
#     instructions = doc.add_paragraph("Instructions:")
#     doc.add_paragraph("1. Attempt all questions")
#     doc.add_paragraph("2. Read each question carefully before answering")
#     doc.add_paragraph("")  # Spacing
    
#     # Add questions with marks
#     for i, (question, difficulty) in enumerate(selected_questions, 1):
#         marks = marks_per_difficulty[difficulty]
#         para = doc.add_paragraph()
#         para.add_run(f"{i}. ").bold = True
#         para.add_run(f"{question} ")
#         para.add_run(f"[{marks} Marks]").italic = True
    
#     # Save document
#     os.makedirs("generated_papers", exist_ok=True)
#     from datetime import datetime
#     current_date = datetime.now().strftime("%Y-%m-%d")
#     file_name = f"generated_papers/{topic}_{current_date}.docx"
#     doc.save(file_name)
#     # file_name = f"generated_papers/{topic}_question_paper.docx"
#     # doc.save(file_name)
    

#     return file_name
import pandas as pd
import random
from docx import Document
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

DATASET_FOLDER = "dataset/"

def get_dataset_for_subject(subject):
    dataset_path = os.path.join(DATASET_FOLDER, f"{subject}.csv")
    if os.path.exists(dataset_path):
        return pd.read_csv(dataset_path)
    return None

def calculate_questions_per_bloom(num_questions, bloom_distribution):
    questions_per_bloom = {}
    remaining_questions = num_questions

    for bloom, percentage in bloom_distribution.items():
        num = int((percentage / 100) * num_questions)
        questions_per_bloom[bloom] = num
        remaining_questions -= num

    if remaining_questions > 0:
        blooms_sorted = sorted(
            bloom_distribution.items(), key=lambda x: x[1], reverse=True
        )
        for i in range(remaining_questions):
            questions_per_bloom[blooms_sorted[i % len(blooms_sorted)][0]] += 1

    return questions_per_bloom

def generate_question_paper(topic, num_questions, total_marks, bloom_distribution):
    df = get_dataset_for_subject(topic)
    if df is None or df.empty:
        raise ValueError(f"No dataset found for topic: {topic}")

    questions_per_bloom = calculate_questions_per_bloom(num_questions, bloom_distribution)

    selected_questions = []

    for bloom, count in questions_per_bloom.items():
        available = df[df['bloom_taxonomy_level'] == bloom]['question'].tolist()
        if available:
            selected = random.sample(available, min(count, len(available)))
            selected_questions.extend([(q, bloom) for q in selected])

    random.shuffle(selected_questions)

    doc = Document()
    header = doc.add_heading(f"Question Paper: {topic}", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    marks_para = doc.add_paragraph(f"Total Marks: {total_marks}")
    marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("Time: 3 Hours")
    doc.add_paragraph("")

    doc.add_paragraph("Instructions:")
    doc.add_paragraph("1. Attempt all questions")
    doc.add_paragraph("2. Read each question carefully before answering")
    doc.add_paragraph("")

    marks_each = total_marks // max(1, len(selected_questions))
    for i, (question, bloom) in enumerate(selected_questions, 1):
        para = doc.add_paragraph()
        para.add_run(f"{i}. ").bold = True
        para.add_run(f"{question} ")
        para.add_run(f"[{marks_each} Marks]").italic = True

    os.makedirs("generated_papers", exist_ok=True)
    current_date = datetime.now().strftime("%Y-%m-%d")
    file_name = f"generated_papers/{topic}_{current_date}.docx"
    doc.save(file_name)

    return file_name
